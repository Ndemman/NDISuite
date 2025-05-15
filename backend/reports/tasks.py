import logging
import os
import json
import tempfile
from datetime import datetime
from celery import shared_task
from django.conf import settings
from ndisuite.utils.llm import get_llm
from langchain.schema import SystemMessage, HumanMessage
from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from .models import Report, OutputField, ReportVersion, ExportedReport

logger = logging.getLogger('ndisuite')


@shared_task
def generate_report_task(report_id):
    """
    Generate report content using AI based on session inputs
    """
    try:
        # Get the report object
        report = Report.objects.get(id=report_id)
        
        # Get template
        template = report.template
        if not template:
            return {
                "success": False,
                "report_id": report_id,
                "error": "No template associated with this report"
            }
        
        # Get session data: files, transcripts
        session = report.session
        
        # Setup LangChain components
        embeddings = OpenAIEmbeddings(
            openai_api_key=settings.OPENAI_API_KEY,
            model=settings.EMBEDDING_MODEL
        )
        
        # Load vector store - use session-specific collection
        vector_store = Chroma(
            collection_name=f"session_{session.id}",
            embedding_function=embeddings,
            persist_directory=settings.VECTOR_STORE_PATH
        )
        
        # Collect vector IDs for files in the current report's session
        file_ids = [str(file.id) for file in session.files.all()]
        
        # Log the file_ids for debugging
        logger.info(f"Files found for session {session.id}: {file_ids}")
        
        # Guard: if no files, skip document RAG
        if not file_ids:
            logger.warning(f"No files found for session {session.id}, document RAG will not be used")
        
        # Document retriever - filtered by file_id and session_id
        doc_retriever = vector_store.as_retriever(
            search_kwargs={
                "k": settings.RAG_TOP_K,
                "filter": {
                    "file_id": {"$in": file_ids},
                    "session_id": str(session.id)  # Add session_id filter for better isolation
                }
                # "mmr": True - Removed (deprecated in LangChain ≥ 0.2.8)
            },
            search_type="mmr"  # Use Maximal Marginal Relevance for diversity
        )
        
        # Transcription retriever - using session-specific collection
        transcript_ns = f"session_{session.id}"
        try:
            # Create a separate vector store for the session's transcripts
            transcript_vs = Chroma(
                collection_name=transcript_ns,
                embedding_function=embeddings,
                persist_directory=settings.VECTOR_STORE_PATH
            )
            
            # Create retriever for transcripts with strict metadata filtering
            transcript_filter = {"session_id": str(session.id), "type": "transcript"}
            trans_retriever = transcript_vs.as_retriever(
                search_kwargs={
                    "k": settings.RAG_TOP_K,
                    "filter": transcript_filter
                },
                search_type="mmr"  # Use MMR instead of deprecated "mmr" parameter
            )
            logger.info(f"Transcript retriever created with filter: {transcript_filter}")
            logger.info(f"Set up transcript retriever for collection: {transcript_ns}")
        except Exception as e:
            logger.warning(f"Could not set up transcript retriever: {str(e)}")
            trans_retriever = None
        
        # Process each field in the template
        template_structure = template.structure
        report_content = {}
        
        for field_name, field_config in template_structure.items():
            # Check if field exists or create it
            field, created = OutputField.objects.get_or_create(
                report=report,
                name=field_name,
                defaults={
                    'label': field_config.get('label', field_name),
                    'field_type': field_config.get('type', 'text'),
                    'options': field_config.get('options', []),
                    'order': field_config.get('order', 0),
                    'validation': field_config.get('validation', {}),
                    'generation_prompt': field_config.get('generation_prompt', '')
                }
            )
            
            # Get generation prompt
            prompt = field.generation_prompt
            if not prompt:
                prompt = f"Generate content for the {field.label} section of an NDIS report."
            
            # Get transcripts
            transcripts = []
            for transcript in session.transcripts.filter(status='completed'):
                try:
                    text = transcript.text
                    if transcript.mongo_id:
                        # Get full text from MongoDB
                        from pymongo import MongoClient
                        client_mongo = MongoClient(settings.MONGODB_URI)
                        db = client_mongo[settings.MONGODB_DB]
                        collection = db['transcripts']
                        doc = collection.find_one({"_id": transcript.mongo_id})
                        if doc:
                            text = doc.get('text', text)
                    
                    transcripts.append(text)
                except Exception as e:
                    logger.error(f"Error getting transcript {transcript.id}: {str(e)}")
            
            # Get relevant context using RAG - dual retrieval approach
            context = ""
            
            # Initialize empty lists for documents
            doc_chunks = []
            trans_chunks = []
            
            # 1. Retrieve document chunks if we have files
            if file_ids:
                try:
                    # Log retrieval attempt
                    logger.info(f"Attempting to retrieve document chunks for prompt: {prompt[:100]}...")
                    doc_chunks = doc_retriever.get_relevant_documents(prompt)
                    logger.info(f"Retrieved {len(doc_chunks)} document chunks with filter: file_ids={file_ids}")
                except Exception as e:
                    logger.error(f"Error during document RAG retrieval: {str(e)}")
            else:
                logger.warning("Skipping document retrieval - no files available")
            
            # 2. Retrieve transcript chunks if transcript retriever is available
            if trans_retriever:
                try:
                    logger.info(f"Attempting to retrieve transcript chunks for prompt: {prompt[:100]}...")
                    trans_chunks = trans_retriever.get_relevant_documents(prompt)
                    logger.info(f"Retrieved {len(trans_chunks)} transcript chunks from collection {transcript_ns}")
                except Exception as e:
                    logger.error(f"Error during transcript RAG retrieval: {str(e)}")
            
            # 3. Merge document and transcript chunks
            all_chunks = doc_chunks + trans_chunks
            
            # 4. If we have chunks, sort by similarity score descending
            if all_chunks:
                try:
                    # Sort by similarity score (if available)
                    all_chunks.sort(key=lambda d: d.metadata.get("similarity", 0), reverse=True)
                    
                    # Optionally trim to RAG_TOP_K
                    all_chunks = all_chunks[:settings.RAG_TOP_K]
                    
                    # Build context from merged chunks
                    for idx, doc in enumerate(all_chunks):
                        source_type = doc.metadata.get('source_type', 'file')
                        if doc.metadata.get('type') == 'transcript':
                            source_type = 'transcript'
                        
                        header = f"[DOC {idx+1} | src={source_type}]"
                        context += header + "\n" + doc.page_content.strip() + "\n\n"
                        
                        # Log metadata for debugging
                        logger.info(f"Using chunk {idx}: {doc.metadata}")
                except Exception as e:
                    logger.error(f"Error merging/sorting chunks: {str(e)}")
            else:
                logger.warning("No chunks retrieved from either documents or transcripts")
                
            # If no chunks were retrieved, use a default message
            if not context.strip():
                context = "No relevant documents or transcripts found for this query."
            
            # Generate content via configured LLM (Gemma/OpenAI)
            llm = get_llm(temperature=0.7)
            system_prompt = f"""You are an expert NDIS report writer. Your task is to generate content for the {field.label} section of an NDIS report.
            
Guidelines:
- Use professional, clear language appropriate for NDIS reports
- Be specific and provide relevant details
- Focus on the client's support needs, goals, and progress
- Avoid vague language and generalizations
- Use person-centered language
- Be objective and evidence-based
- Follow any specific instructions in the prompt

Use the following context to inform your response:
{context}
            """
            
            try:
                response = llm.invoke([
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=prompt)
                ])
                
                # Save the generated content
                generated_content = getattr(response, "content", str(response))
                field.value = generated_content
                field.save()
                
                # Add to report content
                report_content[field_name] = generated_content
            
            except Exception as e:
                logger.error(f"Error generating content for field {field.name}: {str(e)}")
                field.value = f"Error generating content: {str(e)}"
                field.save()
                report_content[field_name] = field.value
        
        # Update report content and status
        report.content = report_content
        report.status = 'generated'
        report.save()
        
        # Create a version
        ReportVersion.objects.create(
            report=report,
            version_number=1,
            content=report_content,
            created_by=report.session.user,
            comment="Initial AI-generated content"
        )
        
        return {
            "success": True,
            "report_id": report_id,
            "status": report.status
        }
    
    except Report.DoesNotExist:
        logger.error(f"Report not found: {report_id}")
        return {
            "success": False,
            "report_id": report_id,
            "error": "Report not found"
        }
    
    except Exception as e:
        logger.error(f"Error generating report {report_id}: {str(e)}")
        
        # Update report status if possible
        try:
            report = Report.objects.get(id=report_id)
            report.save()
        except:
            pass
        
        return {
            "success": False,
            "report_id": report_id,
            "error": str(e)
        }


@shared_task
def export_report_task(report_id, format='pdf'):
    """
    Export a report to PDF or DOCX format
    """
    try:
        # Get the report object
        report = Report.objects.get(id=report_id)
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(suffix=f'.{format}', delete=False) as temp_file:
            temp_file_path = temp_file.name
        
        # Generate file based on format
        if format.lower() == 'pdf':
            success = _generate_pdf(report, temp_file_path)
        elif format.lower() == 'docx':
            success = _generate_docx(report, temp_file_path)
        else:
            raise ValueError(f"Unsupported format: {format}")
        
        if not success:
            raise Exception(f"Failed to generate {format.upper()} file")
        
        # Create exported report record
        from django.core.files.base import File
        
        export = ExportedReport(
            report=report,
            format=format.lower(),
            created_by=report.session.user
        )
        
        # Save file to model
        with open(temp_file_path, 'rb') as f:
            filename = f"{report.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{format}"
            export.file.save(filename, File(f))
        
        # Delete temporary file
        try:
            os.unlink(temp_file_path)
        except:
            pass
        
        return {
            "success": True,
            "report_id": report_id,
            "export_id": str(export.id),
            "format": format,
            "file_url": export.file.url
        }
    
    except Report.DoesNotExist:
        logger.error(f"Report not found: {report_id}")
        return {
            "success": False,
            "report_id": report_id,
            "error": "Report not found"
        }
    
    except Exception as e:
        logger.error(f"Error exporting report {report_id} to {format}: {str(e)}")
        
        # Clean up temporary file if it exists
        try:
            os.unlink(temp_file_path)
        except:
            pass
        
        return {
            "success": False,
            "report_id": report_id,
            "error": str(e)
        }


def _generate_pdf(report, output_path):
    """
    Generate a PDF file for the report
    """
    try:
        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'Title', 
            parent=styles['Title'],
            alignment=1,  # Center
            spaceAfter=12
        )
        
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            spaceAfter=6,
            spaceBefore=12
        )
        
        content_style = ParagraphStyle(
            'Content',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            spaceBefore=0,
            spaceAfter=10
        )
        
        # Create document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Prepare content elements
        elements = []
        
        # Add title
        elements.append(Paragraph(report.title, title_style))
        elements.append(Spacer(1, 12))
        
        # Add date
        elements.append(Paragraph(f"Generated: {datetime.now().strftime('%d %B %Y')}", styles['Normal']))
        elements.append(Spacer(1, 24))
        
        # Add fields
        fields = report.fields.all().order_by('order')
        for field in fields:
            # Add field heading
            elements.append(Paragraph(field.label, heading_style))
            
            # Add field content
            elements.append(Paragraph(field.value, content_style))
        
        # Build document
        doc.build(elements)
        
        return True
    
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        return False


def _generate_docx(report, output_path):
    """
    Generate a DOCX file for the report
    """
    try:
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading(report.title, level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y')}")
        date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Add fields
        fields = report.fields.all().order_by('order')
        for field in fields:
            # Add field heading
            heading = doc.add_heading(field.label, level=2)
            
            # Add field content
            doc.add_paragraph(field.value)
        
        # Save document
        doc.save(output_path)
        
        return True
    
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        return False


@shared_task
def refine_field_content_task(field_id, instruction, tone='professional', length='maintain', format_style='paragraph'):
    """
    Refine field content using AI
    """
    try:
        # Get the field object
        field = OutputField.objects.get(id=field_id)
        
        # Prepare the prompt
        system_prompt = f"""You are an expert NDIS report editor. Your task is to refine the content based on the user's instructions.
        
Guidelines:
- Tone: {tone} (e.g., professional, conversational, formal, etc.)
- Length: {length} (maintain, shorter, longer)
- Format: {format_style} (paragraph, bullet points, numbered list)
- Maintain person-centered language
- Be specific and provide relevant details
- Focus on the client's support needs, goals, and progress
- Use professional language appropriate for NDIS reports

Instruction from the user: {instruction}

Please refine the original content according to these instructions while maintaining professional language suitable for an NDIS report.
"""
        
        # Generate refined content via configured LLM (Gemma/OpenAI)
        llm = get_llm(temperature=0.7)
        response = llm.invoke([
            SystemMessage(content=system_prompt),
            HumanMessage(content=field.value)
        ])
        
        # Extract content (LangChain returns ChatMessage)
        refined_content = getattr(response, "content", str(response))
        
        # Update the field content
        field.value = refined_content
        field.save()
        
        # Update the report content
        report = field.report
        report_content = report.content
        report_content[field.name] = refined_content
        report.content = report_content
        
        # Update report status
        if report.status == 'generated':
            report.status = 'refined'
        
        report.save()
        
        return {
            "success": True,
            "field_id": field_id,
            "report_id": str(report.id)
        }
    
    except OutputField.DoesNotExist:
        logger.error(f"Field not found: {field_id}")
        return {
            "success": False,
            "field_id": field_id,
            "error": "Field not found"
        }
    
    except Exception as e:
        logger.error(f"Error refining field {field_id}: {str(e)}")
        return {
            "success": False,
            "field_id": field_id,
            "error": str(e)
        }
